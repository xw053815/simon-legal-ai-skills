"""测试 preserve_format_revisions 的格式保留能力。"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from preserve_format_revisions import apply_revisions_to_doc


def make_test_doc(path: Path) -> None:
    d = Document()
    p1 = d.add_paragraph()
    p1.add_run("甲方应在").font.name = "SimSun"
    p1.add_run("30").font.name = "Times New Roman"
    p1.add_run("日内付款。").font.name = "SimSun"

    p2 = d.add_paragraph()
    p2.add_run("乙方负责交付").font.name = "SimSun"
    p2.add_run("产品").font.name = "SimSun"
    p2.add_run("。").font.name = "SimSun"

    d.save(path)
    print(f"已生成测试文件：{path}")


def inspect_doc(path: Path, label: str) -> None:
    doc = Document(path)
    print(f"\n=== {label} ===")
    for i, para in enumerate(doc.paragraphs):
        print(f"Paragraph {i}:")
        for j, run in enumerate(para.runs):
            print(f"  run {j}: {run.text!r} | font={run.font.name} | bold={run.bold}")

        # 也检查修订标记
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        ins = para._p.findall(f'.//{ns}ins')
        dels = para._p.findall(f'.//{ns}del')
        if ins or dels:
            print(f"  修订标记: {len(ins)} 插入, {len(dels)} 删除")
            for e in ins:
                author = e.get(f'{ns}author')
                text = ''.join(t.text or '' for t in e.iter(f'{ns}t'))
                print(f"    插入: {text!r} by {author}")
            for e in dels:
                author = e.get(f'{ns}author')
                text = ''.join(t.text or '' for t in e.iter(f'{ns}delText'))
                print(f"    删除: {text!r} by {author}")


if __name__ == "__main__":
    base = Path(r"[用户目录]\WorkBuddy\[项目目录]\02_scratch")
    src = base / "preserve_format_test.docx"
    dst = base / "preserve_format_test_revised.docx"

    make_test_doc(src)
    inspect_doc(src, "原文件")

    doc = Document(src)
    log = apply_revisions_to_doc(
        doc,
        replacements=[("30", "60")],
        insertions_after=[("60", "（宽限期）")],
        deletions=["产品"],
        author="律师",
    )
    doc.save(dst)

    print("\n操作日志:")
    for line in log:
        print(f"  {line}")

    inspect_doc(dst, "修订后")
