"""preserve_format_revisions.py — 在保留原 run 格式的前提下生成 Word/WPS 修订痕迹。

本模块不依赖 docx-editor 的文本替换逻辑，而是直接基于 python-docx 的 XML 元素操作，
确保 <w:ins> / <w:del> 内的文本继承其所在位置原 run 的 rPr（字体、字号、加粗、斜体、
下划线、颜色等），避免修订后格式与原文档不匹配。
"""
from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Iterable

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run


# 全局计数器，用于生成唯一的 w:id
def _revision_id_generator() -> Iterable[int]:
    """生成递增的修订 ID。"""
    i = 1
    while True:
        yield i
        i += 1


_revision_ids = _revision_id_generator()


def _next_revision_id() -> int:
    return next(_revision_ids)


def _now_iso() -> str:
    """返回当前时间的 ISO 8601 UTC 字符串（Word 习惯）。"""
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _set_xml_space(t_elem: OxmlElement, text: str) -> None:
    """如果文本首尾有空格，设置 xml:space="preserve"。"""
    if text and (text[0].isspace() or text[-1].isspace()):
        t_elem.set(qn("xml:space"), "preserve")


def _make_text_element(text: str, is_deletion: bool = False) -> OxmlElement:
    """创建 <w:t> 或 <w:delText> 元素。"""
    tag = "w:delText" if is_deletion else "w:t"
    t = OxmlElement(tag)
    t.text = text
    _set_xml_space(t, text)
    return t


def _clone_rPr(run: Run) -> OxmlElement | None:
    """复制 run 的 rPr（格式属性）。如果 run 没有 rPr，返回 None。"""
    r = run._r
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        return None
    return deepcopy(rPr)


def _make_run_element(text: str, rPr: OxmlElement | None, is_deletion: bool = False) -> OxmlElement:
    """创建一个 <w:r> 元素，包含可选的 rPr 和文本。"""
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(deepcopy(rPr))
    r.append(_make_text_element(text, is_deletion))
    return r


def _make_revision_element(
    rev_type: str, text: str, rPr: OxmlElement | None, author: str, date: str | None = None
) -> OxmlElement:
    """创建 <w:ins> 或 <w:del> 元素，内部包含一个 <w:r>。"""
    elem = OxmlElement(f"w:{rev_type}")
    elem.set(qn("w:id"), str(_next_revision_id()))
    elem.set(qn("w:author"), author)
    elem.set(qn("w:date"), date or _now_iso())
    elem.append(_make_run_element(text, rPr, is_deletion=(rev_type == "del")))
    return elem


def _find_nth(text: str, substring: str, occurrence: int = 0) -> int:
    """查找子字符串第 occurrence 次出现的位置（0-based）。未找到返回 -1。"""
    if not substring:
        return -1
    start = 0
    for _ in range(occurrence + 1):
        pos = text.find(substring, start)
        if pos == -1:
            return -1
        start = pos + 1
    return start - 1


@dataclass(frozen=True)
class CharPosition:
    """段落内一个字符的位置。"""
    char: str
    node: OxmlElement  # 文本节点 <w:t> 或 <w:delText>
    offset: int        # 在节点内的偏移
    run: OxmlElement   # 父 <w:r> 元素
    is_inside_ins: bool


class ParagraphIndex:
    """把段落内所有可见文本（包括 <w:ins> 内，排除 <w:del> 内）展开成字符索引表。"""

    def __init__(self, paragraph: Paragraph):
        self.positions: list[CharPosition] = []
        self._walk(paragraph._p, inside_ins=False)
        self.text = "".join(pos.char for pos in self.positions)

    def _walk(self, elem: OxmlElement, *, inside_ins: bool) -> None:
        """递归遍历段落 XML。"""
        for child in elem:
            tag = child.tag
            if tag == qn("w:del"):
                continue
            if tag == qn("w:ins"):
                self._walk(child, inside_ins=True)
                continue
            if tag in (qn("w:t"), qn("w:delText")):
                text = child.text or ""
                parent = child.getparent()
                run = parent if parent is not None and parent.tag == qn("w:r") else None
                for offset, char in enumerate(text):
                    self.positions.append(CharPosition(char, child, offset, run, inside_ins))
                continue
            # 递归处理其它容器（如 <w:r>、<w:smartTag> 等）
            self._walk(child, inside_ins=inside_ins)

    def find(self, target: str, occurrence: int = 0) -> tuple[int, int] | None:
        """返回目标文本在段落中的 [start, end) 字符索引范围。未找到返回 None。"""
        start = _find_nth(self.text, target, occurrence)
        if start == -1:
            return None
        return (start, start + len(target))

    def position(self, idx: int) -> CharPosition:
        return self.positions[idx]


def _split_run_text(run: Run, split_offsets: list[int]) -> list[str]:
    """按偏移量把 run 的文本分割成多个片段。"""
    text = run.text
    parts = []
    prev = 0
    for offset in sorted(split_offsets):
        if offset > prev:
            parts.append(text[prev:offset])
        prev = offset
    if prev < len(text):
        parts.append(text[prev:])
    return parts


def _replace_single_run(
    paragraph: Paragraph,
    run: OxmlElement,
    old_text: str,
    new_text: str,
    offset_in_run: int,
    author: str,
    date: str | None = None,
) -> None:
    """在单个 run 内完成替换，并生成 <w:del> 和 <w:ins>。"""
    rPr = _clone_rPr_from_elem(run)
    t_node = run.find(qn("w:t"))
    if t_node is None:
        raise ValueError("run 内没有 w:t 节点")
    full_text = t_node.text or ""

    before = full_text[:offset_in_run]
    after = full_text[offset_in_run + len(old_text) :]

    p = paragraph._p
    old_r = run

    # 插入顺序：before -> del -> ins -> after，全部在 old_r 之前
    if before:
        p.insert(p.index(old_r), _make_run_element(before, rPr))
    if old_text:
        p.insert(p.index(old_r), _make_revision_element("del", old_text, rPr, author, date))
    if new_text:
        p.insert(p.index(old_r), _make_revision_element("ins", new_text, rPr, author, date))
    if after:
        p.insert(p.index(old_r), _make_run_element(after, rPr))

    p.remove(old_r)


def _replace_across_runs(
    paragraph: Paragraph,
    positions: list[CharPosition],
    old_text: str,
    new_text: str,
    author: str,
    date: str | None = None,
) -> None:
    """处理目标文本跨多个 run 的替换，保留每个 run 的格式。"""
    start_pos = positions[0]
    end_pos = positions[-1]
    first_run = start_pos.run
    last_run = end_pos.run

    # 收集所有受影响的 run（按文档顺序）
    unique_runs: list[OxmlElement] = []
    seen = set()
    for pos in positions:
        if pos.run is not None and id(pos.run) not in seen:
            seen.add(id(pos.run))
            unique_runs.append(pos.run)

    if first_run is None or last_run is None:
        raise ValueError("跨 run 替换时找不到有效 run")

    first_rPr = _clone_rPr_from_elem(first_run)
    last_rPr = _clone_rPr_from_elem(last_run)

    first_t = first_run.find(qn("w:t"))
    first_text = first_t.text or "" if first_t is not None else ""
    first_before = first_text[: start_pos.offset]
    first_matched = first_text[start_pos.offset :]

    last_t = last_run.find(qn("w:t"))
    last_text = last_t.text or "" if last_t is not None else ""
    last_matched = last_text[: end_pos.offset + 1]
    last_after = last_text[end_pos.offset + 1 :]

    p = paragraph._p
    anchor = first_run

    if first_before:
        p.insert(p.index(anchor), _make_run_element(first_before, first_rPr))
    if first_matched:
        p.insert(
            p.index(anchor),
            _make_revision_element("del", first_matched, first_rPr, author, date),
        )

    # 中间 runs：整体删除
    for r in unique_runs[1:-1]:
        rPr = _clone_rPr_from_elem(r)
        t = r.find(qn("w:t"))
        text = t.text or "" if t is not None else ""
        p.insert(p.index(anchor), _make_revision_element("del", text, rPr, author, date))

    if last_matched and len(unique_runs) > 1:
        p.insert(
            p.index(anchor),
            _make_revision_element("del", last_matched, last_rPr, author, date),
        )

    if new_text:
        # 新文本使用第一个 run 的格式
        p.insert(p.index(anchor), _make_revision_element("ins", new_text, first_rPr, author, date))

    if last_after and len(unique_runs) > 1:
        p.insert(p.index(anchor), _make_run_element(last_after, last_rPr))

    # 删除所有旧 run
    for r in unique_runs:
        if r.getparent() is not None:
            r.getparent().remove(r)


def replace_in_paragraph(
    paragraph: Paragraph,
    old_text: str,
    new_text: str,
    *,
    author: str,
    occurrence: int = 0,
    date: str | None = None,
) -> bool:
    """在段落中替换文本，并生成修订痕迹，同时保留原格式。

    支持单 run 内替换和跨 run 替换。目标文本在 <w:ins> 内时暂不支持（避免嵌套）。
    返回 True 表示成功，False 表示未找到文本。
    """
    index = ParagraphIndex(paragraph)
    range_ = index.find(old_text, occurrence)
    if range_ is None:
        return False

    start, end = range_
    positions = index.positions[start:end]

    if any(pos.is_inside_ins for pos in positions):
        # 不支持在已有 <w:ins> 内再修改（避免嵌套）
        return False

    runs = {pos.run for pos in positions if pos.run is not None}
    if len(runs) == 1:
        run = positions[0].run
        offset = positions[0].offset
        _replace_single_run(paragraph, run, old_text, new_text, offset, author, date)
    else:
        _replace_across_runs(paragraph, positions, old_text, new_text, author, date)
    return True


def delete_in_paragraph(
    paragraph: Paragraph,
    text: str,
    *,
    author: str,
    occurrence: int = 0,
    date: str | None = None,
) -> bool:
    """在段落中删除文本，并生成 <w:del> 修订痕迹，保留原格式。"""
    return replace_in_paragraph(
        paragraph, text, "", author=author, occurrence=occurrence, date=date
    )


def insert_after_in_paragraph(
    paragraph: Paragraph,
    anchor: str,
    text: str,
    *,
    author: str,
    occurrence: int = 0,
    date: str | None = None,
) -> bool:
    """在锚点文本之后插入内容，并生成 <w:ins> 修订痕迹，保留原格式。"""
    index = ParagraphIndex(paragraph)
    range_ = index.find(anchor, occurrence)
    if range_ is None:
        return False

    anchor_end = range_[1] - 1
    pos = index.position(anchor_end)
    run = pos.run
    if run is None:
        return False

    rPr = _clone_rPr_from_elem(run)
    t_node = run.find(qn("w:t"))
    if t_node is None:
        return False
    full_text = t_node.text or ""

    # 如果锚点在 <w:ins> 内，在外层 ins 后追加新的 <w:ins>，避免嵌套
    if pos.is_inside_ins:
        ins_ancestor = _find_ancestor_ins(run)
        if ins_ancestor is None:
            return False
        p = paragraph._p
        ins_new = _make_revision_element("ins", text, rPr, author, date)
        p.insert(p.index(ins_ancestor) + 1, ins_new)
        return True

    offset = pos.offset + 1
    before = full_text[:offset]
    after = full_text[offset:]

    p = paragraph._p
    old_r = run

    if before:
        p.insert(p.index(old_r), _make_run_element(before, rPr))
    p.insert(p.index(old_r), _make_revision_element("ins", text, rPr, author, date))
    if after:
        p.insert(p.index(old_r), _make_run_element(after, rPr))
    p.remove(old_r)
    return True


def insert_before_in_paragraph(
    paragraph: Paragraph,
    anchor: str,
    text: str,
    *,
    author: str,
    occurrence: int = 0,
    date: str | None = None,
) -> bool:
    """在锚点文本之前插入内容，并生成 <w:ins> 修订痕迹，保留原格式。"""
    index = ParagraphIndex(paragraph)
    range_ = index.find(anchor, occurrence)
    if range_ is None:
        return False

    start = range_[0]
    pos = index.position(start)
    run = pos.run
    if run is None:
        return False

    rPr = _clone_rPr_from_elem(run)
    t_node = run.find(qn("w:t"))
    if t_node is None:
        return False
    full_text = t_node.text or ""

    # 如果锚点在 <w:ins> 内，在外层 ins 前插入新的 <w:ins>
    if pos.is_inside_ins:
        ins_ancestor = _find_ancestor_ins(run)
        if ins_ancestor is None:
            return False
        p = paragraph._p
        ins_new = _make_revision_element("ins", text, rPr, author, date)
        p.insert(p.index(ins_ancestor), ins_new)
        return True

    offset = pos.offset
    before = full_text[:offset]
    after = full_text[offset:]

    p = paragraph._p
    old_r = run

    if before:
        p.insert(p.index(old_r), _make_run_element(before, rPr))
    p.insert(p.index(old_r), _make_revision_element("ins", text, rPr, author, date))
    if after:
        p.insert(p.index(old_r), _make_run_element(after, rPr))
    p.remove(old_r)
    return True


def _clone_rPr_from_elem(run: OxmlElement) -> OxmlElement | None:
    """从 <w:r> 元素复制 rPr。"""
    rPr = run.find(qn("w:rPr"))
    if rPr is None:
        return None
    return deepcopy(rPr)


def _find_ancestor_ins(elem: OxmlElement) -> OxmlElement | None:
    """向上查找最近的 <w:ins> 祖先。"""
    p = elem.getparent()
    while p is not None:
        if p.tag == qn("w:ins"):
            return p
        p = p.getparent()
    return None


def _iter_all_paragraphs(doc):
    """遍历文档中所有段落，包括表格单元格内的段落。"""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def apply_revisions_to_doc(
    doc,
    replacements: list[tuple[str, str]] | None = None,
    insertions_after: list[tuple[str, str]] | None = None,
    insertions_before: list[tuple[str, str]] | None = None,
    deletions: list[str] | list[tuple[str]] | None = None,
    *,
    author: str = "律师",
    date: str | None = None,
    stop_on_error: bool = True,
) -> list[str]:
    """对文档应用一批修订操作，返回执行日志。"""
    replacements = list(replacements or [])
    insertions_after = list(insertions_after or [])
    insertions_before = list(insertions_before or [])
    deletions = list(deletions or [])
    # 归一化：允许 deletions 元素为 str 或 (str,)
    normalized_deletions: list[str] = []
    for item in deletions:
        if isinstance(item, (tuple, list)):
            normalized_deletions.extend(str(x) for x in item if x)
        else:
            normalized_deletions.append(str(item))

    log: list[str] = []

    for old, new in replacements:
        found = False
        for para in _iter_all_paragraphs(doc):
            if replace_in_paragraph(para, old, new, author=author, date=date):
                log.append(f"替换: {old!r} -> {new!r}")
                found = True
                break
        if not found:
            msg = f"未找到替换目标: {old!r}"
            log.append(msg)
            if stop_on_error:
                raise TextNotFoundError(msg)

    for anchor, text in insertions_after:
        found = False
        for para in _iter_all_paragraphs(doc):
            if insert_after_in_paragraph(para, anchor, text, author=author, date=date):
                log.append(f"插入后: {anchor!r} -> +{text!r}")
                found = True
                break
        if not found:
            msg = f"未找到插入锚点: {anchor!r}"
            log.append(msg)
            if stop_on_error:
                raise TextNotFoundError(msg)

    for anchor, text in insertions_before:
        found = False
        for para in _iter_all_paragraphs(doc):
            if insert_before_in_paragraph(para, anchor, text, author=author, date=date):
                log.append(f"插入前: {anchor!r} -> +{text!r}")
                found = True
                break
        if not found:
            msg = f"未找到插入锚点: {anchor!r}"
            log.append(msg)
            if stop_on_error:
                raise TextNotFoundError(msg)

    for text in normalized_deletions:
        found = False
        for para in _iter_all_paragraphs(doc):
            if delete_in_paragraph(para, text, author=author, date=date):
                log.append(f"删除: {text!r}")
                found = True
                break
        if not found:
            msg = f"未找到删除目标: {text!r}"
            log.append(msg)
            if stop_on_error:
                raise TextNotFoundError(msg)

    return log


class TextNotFoundError(Exception):
    """目标文本未找到。"""
