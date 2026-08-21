# -*- coding: utf-8 -*-
"""audit_docx.py — 用 python-docx 读取 OfficeCLI 产物，交叉核验格式参数

用法：
    report = audit_docx('C:/path/to/output.docx')
    print_report(report)

集成到四阶段任务工作流：
    I 阶段 → OfficeCLI 执行，生成 V1 产物
    V 阶段 → 1. OfficeCLI validate + view issues
             2. python-docx audit_docx() + print_report()
             3. 比对两份报告，确认无差异项
    修复   → 任一引擎发现问题 → 用 OfficeCLI 修复 → 重新走双核验

禁止事项：
    1. 禁止在审查时修改文件（python-docx 审查脚本只读不写，不调用 doc.save()）
    2. 禁止跳过任何一方（OfficeCLI validate 通过不代表排版正确）
    3. 禁止将审查脚本的输出作为最终交付物（审查报告是内部的，不推送给客户）
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Twips
import re


def audit_docx(path):
    """返回格式检查报告"""
    doc = Document(path)
    report = {
        'page_margins': None,
        'font_coverage': None,
        'line_spacing_issues': [],
        'first_line_indent_issues': [],
        'font_size_issues': [],
        'east_asia_issues': [],
    }

    # 1. 页边距检查
    sec = doc.sections[0]
    report['page_margins'] = {
        'top_cm': round(sec.top_margin / 360000, 2),
        'bottom_cm': round(sec.bottom_margin / 360000, 2),
        'left_cm': round(sec.left_margin / 360000, 2),
        'right_cm': round(sec.right_margin / 360000, 2),
    }

    # 2. 四层字体覆盖率
    total = 0
    ok = 0
    for p in doc.paragraphs:
        for run in p.runs:
            total += 1
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None and rf.get(qn('w:eastAsia')):
                    ok += 1
                    # 检查 eastAsia 是否被污染为西文字体
                    ascii_font = rf.get(qn('w:ascii'), '')
                    if any(cn in ascii_font for cn in ['宋体', '仿宋', '黑体', '楷体', 'SimSun', 'FangSong', 'SimHei', 'KaiTi']):
                        report['east_asia_issues'].append(f'ascii 被中文字体污染: {ascii_font}')

    # 表格内
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        total += 1
                        rPr = run._element.find(qn('w:rPr'))
                        if rPr is not None:
                            rf = rPr.find(qn('w:rFonts'))
                            if rf is not None and rf.get(qn('w:eastAsia')):
                                ok += 1

    # 页眉页脚
    for section in doc.sections:
        for container in [section.header, section.footer]:
            for p in container.paragraphs:
                for run in p.runs:
                    total += 1
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        rf = rPr.find(qn('w:rFonts'))
                        if rf is not None and rf.get(qn('w:eastAsia')):
                            ok += 1

    report['font_coverage'] = f'{ok}/{total}'

    # 3. 行距检查（取前 5 个段落采样）
    for i, p in enumerate(doc.paragraphs[:5]):
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                line = spacing.get(qn('w:line'))
                lineRule = spacing.get(qn('w:lineRule'))
                report['line_spacing_issues'].append({
                    'para': i,
                    'line': line,
                    'lineRule': lineRule,
                    'text': p.text[:30] if p.text else ''
                })

    # 4. 首行缩进检查（取前 10 个段落）
    for i, p in enumerate(doc.paragraphs[:10]):
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                firstLine = ind.get(qn('w:firstLine'))
                firstLineChars = ind.get(qn('w:firstLineChars'))
                if firstLineChars and int(firstLineChars) != 200:
                    report['first_line_indent_issues'].append({
                        'para': i,
                        'firstLineChars': firstLineChars,
                        'text': p.text[:30] if p.text else ''
                    })

    # 5. 字号检查（取前 10 个段落采样）
    for i, p in enumerate(doc.paragraphs[:10]):
        for run in p.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    val = sz.get(qn('w:val'))
                    pt = int(val) / 2
                    if pt not in [9, 10.5, 12, 15, 18, 22]:
                        report['font_size_issues'].append({
                            'para': i,
                            'sz_val': val,
                            'pt': pt,
                            'text': run.text[:20] if run.text else ''
                        })

    return report


def print_report(report):
    """格式化为可读报告"""
    lines = []
    lines.append('=' * 50)
    lines.append('  双引擎交叉核验报告')
    lines.append('=' * 50)

    # 页边距
    m = report['page_margins']
    expected = '法院版:上3.7 下3.5 左2.8 右2.8 | 客户版:上2.2 下2.2 左2.4 右2.4'
    actual = f'上{m["top_cm"]} 下{m["bottom_cm"]} 左{m["left_cm"]} 右{m["right_cm"]}'
    lines.append(f'\n📐 页边距: {actual}')
    lines.append(f'   预期: {expected}')
    lines.append(f'   {"✅" if abs(m["top_cm"]-3.7)<0.1 or abs(m["top_cm"]-2.2)<0.1 else "❌"} 核对（任一边距档位匹配即通过）')

    # 字体覆盖率
    lines.append(f'\n🔤 字体覆盖率: {report["font_coverage"]}')
    if report['east_asia_issues']:
        for issue in report['east_asia_issues']:
            lines.append(f'   ❌ {issue}')
    else:
        lines.append(f'   ✅ 无 eastAsia 污染')

    # 行距
    lines.append(f'\n📏 行距检查（采样前5段）:')
    for item in report['line_spacing_issues']:
        lines.append(f'   段{item["para"]}: line={item["line"]} rule={item["lineRule"]} "{item["text"]}"')

    # 首行缩进
    if report['first_line_indent_issues']:
        lines.append(f'\n⚠️ 首行缩进异常:')
        for item in report['first_line_indent_issues']:
            lines.append(f'   段{item["para"]}: firstLineChars={item["firstLineChars"]} "{item["text"]}"')
    else:
        lines.append(f'\n✅ 首行缩进: 全部正确')

    # 字号异常
    if report['font_size_issues']:
        lines.append(f'\n⚠️ 字号异常:')
        for item in report['font_size_issues']:
            lines.append(f'   段{item["para"]}: {item["pt"]}pt (sz={item["sz_val"]}) "{item["text"]}"')
    else:
        lines.append(f'\n✅ 字号: 全部在预期范围内')

    print('\n'.join(lines))
    return '\n'.join(lines)


# 使用示例
# report = audit_docx('C:/path/to/output.docx')
# print_report(report)
