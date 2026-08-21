# -*- coding: utf-8 -*-
"""python-docx 字体修复：遍历所有 run（正文+表格+页眉页脚），确保每个 run 有完整的 rFonts

覆盖四层（lesson learned 2026-07-10）：
① doc.paragraphs           — 正文段落
② doc.tables               — 表格内单元格（最易遗漏！三度修复失败均因此）
③ section.header.paragraphs — 页眉
④ section.footer.paragraphs — 页脚

使用顺序（强制）：
    set_doc_defaults_font(doc)   # 第一步：设兜底层
    ... 创建内容 ...
    fix_all_fonts(doc)           # 第二步：遍历所有 run 补全 rFonts
    doc.save('output.docx')
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_doc_defaults_font(doc):
    """设置文档默认字体，作为所有 run 的兜底层"""
    style_element = doc.styles.element
    docDefaults = style_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        style_element.insert(0, docDefaults)

    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)

    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    print('  已设置文档默认字体：仿宋 + Times New Roman')


def fix_all_fonts(doc):
    """遍历文档所有 run，确保每个 run 有完整的 rFonts"""
    count = 0

    # ① 正文段落
    for p in doc.paragraphs:
        for run in p.runs:
            rPr = run._element.get_or_add_rPr()
            rf = rPr.find(qn('w:rFonts'))
            if rf is None:
                rf = OxmlElement('w:rFonts')
                rPr.insert(0, rf)
            rf.set(qn('w:ascii'), 'Times New Roman')
            rf.set(qn('w:hAnsi'), 'Times New Roman')
            rf.set(qn('w:eastAsia'), '仿宋')
            count += 1

    # ② 表格内单元格（最易遗漏！）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        rPr = run._element.get_or_add_rPr()
                        rf = rPr.find(qn('w:rFonts'))
                        if rf is None:
                            rf = OxmlElement('w:rFonts')
                            rPr.insert(0, rf)
                        rf.set(qn('w:ascii'), 'Times New Roman')
                        rf.set(qn('w:hAnsi'), 'Times New Roman')
                        rf.set(qn('w:eastAsia'), '仿宋')
                        count += 1

    # ③④ 页眉页脚
    for section in doc.sections:
        for container in [section.header, section.footer]:
            for p in container.paragraphs:
                for run in p.runs:
                    rPr = run._element.get_or_add_rPr()
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is None:
                        rf = OxmlElement('w:rFonts')
                        rPr.insert(0, rf)
                    rf.set(qn('w:ascii'), 'Times New Roman')
                    rf.set(qn('w:hAnsi'), 'Times New Roman')
                    rf.set(qn('w:eastAsia'), '宋体')
                    count += 1

    print(f'  已修复 {count} 个 run 的字体设置')
    return count


# 四层防御体系（2026-07-10 经验总结）：
# | 层级 | 覆盖范围 | 忘了会怎样 |
# |------|---------|-----------|
# | ① docDefaults | 文档级默认字体（兜底） | 所有无 rFonts 的 run 回退到 MS 明朝 |
# | ② doc.paragraphs | 正文段落 | 正文中文显示为 MS 明朝 |
# | ③ doc.tables | **表格内单元格** ← 最易遗漏 | 表格内中文显示为 MS 明朝 |
# | ④ sections.headers/footers | 页眉页脚 | 页眉页脚显示异常 |
#
# 斯里兰卡项目教训：fix_all_fonts() 只覆盖了①②④，遗漏③表格。
# 三轮修复均失败，第四轮才补全。此后任何排版修复必须四层全覆盖。
